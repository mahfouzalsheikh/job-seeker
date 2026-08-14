from __future__ import annotations

import hashlib
import re
import shutil
import subprocess
import tempfile
import zipfile
from datetime import timedelta
from html.parser import HTMLParser
from pathlib import Path
from typing import Any
from xml.etree import ElementTree

from django.db import transaction
from django.db.models import Count, Q
from django.utils import timezone

from .ai import (
    chunk_text,
    clean_text,
    cosine_similarity,
    detect_skills,
    analyze_candidate_resume,
    extract_job,
    extract_profile_facts,
    keywords,
    stable_hash,
    tailor_resume,
)
from .models import (
    Application,
    Artifact,
    CandidateProfile,
    JobMatch,
    JobPosting,
    ProfileChunk,
    ProfileDocument,
    ProfileFact,
    Resume,
)


class _ReadableHTMLParser(HTMLParser):
    def __init__(self):
        super().__init__()
        self.parts: list[str] = []
        self.ignored_depth = 0

    def handle_starttag(self, tag, attrs):
        if tag in {'script', 'style', 'noscript'}:
            self.ignored_depth += 1
        elif tag in {'p', 'div', 'li', 'br', 'h1', 'h2', 'h3', 'h4', 'tr'}:
            self.parts.append('\n')

    def handle_endtag(self, tag):
        if tag in {'script', 'style', 'noscript'} and self.ignored_depth:
            self.ignored_depth -= 1
        elif tag in {'p', 'div', 'li', 'h1', 'h2', 'h3', 'h4', 'tr'}:
            self.parts.append('\n')

    def handle_data(self, data):
        if not self.ignored_depth:
            self.parts.append(data)


def _extract_html(path: str) -> str:
    parser = _ReadableHTMLParser()
    with open(path, 'r', encoding='utf-8', errors='ignore') as handle:
        parser.feed(handle.read())
    return '\n'.join(line.strip() for line in ''.join(parser.parts).splitlines() if line.strip())


def _extract_rtf(path: str) -> str:
    with open(path, 'r', encoding='utf-8', errors='ignore') as handle:
        content = handle.read()
    content = re.sub(r"\\'([0-9a-fA-F]{2})", lambda match: bytes.fromhex(match.group(1)).decode('latin-1'), content)
    content = re.sub(r'\\par[d]?\s*', '\n', content)
    content = re.sub(r'\\[a-zA-Z]+-?\d*\s?', '', content)
    return re.sub(r'[{}]', '', content)


def _extract_odt(path: str) -> str:
    with zipfile.ZipFile(path) as archive:
        root = ElementTree.fromstring(archive.read('content.xml'))
    return '\n'.join(value.strip() for value in root.itertext() if value.strip())


def _ocr_pdf(path: str) -> str:
    if not shutil.which('pdftoppm') or not shutil.which('tesseract'):
        return ''
    with tempfile.TemporaryDirectory(prefix='forth-resume-ocr-') as temp_dir:
        prefix = str(Path(temp_dir) / 'page')
        subprocess.run(
            ['pdftoppm', '-f', '1', '-l', '12', '-r', '180', '-png', path, prefix],
            check=True,
            capture_output=True,
            timeout=90,
        )
        pages = []
        for image_path in sorted(Path(temp_dir).glob('page-*.png')):
            result = subprocess.run(
                ['tesseract', str(image_path), 'stdout'],
                check=True,
                capture_output=True,
                text=True,
                timeout=45,
            )
            pages.append(result.stdout)
        return '\n\n'.join(pages)


def extract_text_from_upload(path: str) -> str:
    suffix = Path(path).suffix.lower()
    try:
        if suffix == '.pdf':
            from pypdf import PdfReader

            reader = PdfReader(path)
            text = '\n\n'.join(page.extract_text() or '' for page in reader.pages)
            return text if clean_text(text) else _ocr_pdf(path)
        if suffix == '.docx':
            from docx import Document

            doc = Document(path)
            paragraphs = [paragraph.text for paragraph in doc.paragraphs]
            for table in doc.tables:
                paragraphs.extend(cell.text for row in table.rows for cell in row.cells)
            return '\n'.join(paragraphs)
        if suffix == '.doc':
            if not shutil.which('antiword'):
                raise ValueError('Legacy .doc reading is unavailable. Upload DOCX, PDF, or HTML instead.')
            result = subprocess.run(['antiword', path], check=True, capture_output=True, text=True, timeout=45)
            return result.stdout
        if suffix in {'.html', '.htm'}:
            return _extract_html(path)
        if suffix == '.rtf':
            return _extract_rtf(path)
        if suffix == '.odt':
            return _extract_odt(path)
        with open(path, 'r', encoding='utf-8', errors='ignore') as handle:
            return handle.read()
    except Exception as exc:
        raise ValueError(f'Could not extract text from upload: {exc}') from exc


@transaction.atomic
def ingest_profile_document(document: ProfileDocument) -> dict[str, Any]:
    document.status = 'processing'
    document.status_message = 'Extracting profile facts'
    document.save(update_fields=['status', 'status_message', 'updated_at'])

    text = document.raw_text
    if document.upload and not text:
        text = extract_text_from_upload(document.upload.path)
        document.raw_text = text
        document.save(update_fields=['raw_text', 'updated_at'])

    cleaned = clean_text(text)
    if not cleaned:
        document.status = 'failed'
        document.status_message = 'No text was found in the document.'
        document.save(update_fields=['status', 'status_message', 'updated_at'])
        return {'created_facts': 0, 'created_chunks': 0}

    ProfileChunk.objects.filter(document=document).delete()
    chunks = []
    for index, chunk in enumerate(chunk_text(text)):
        chunks.append(ProfileChunk.objects.create(
            owner=document.owner,
            document=document,
            text=chunk,
            token_count=max(1, len(chunk.split())),
            metadata={'index': index},
        ))

    extracted = (
        analyze_candidate_resume(text, title=document.title)
        if document.kind == 'resume'
        else extract_profile_facts(text, title=document.title)
    )
    existing_keys = set(
        ProfileFact.objects.filter(owner=document.owner)
        .values_list('fact_type', 'title', 'statement')
    )
    created = 0
    ambiguity_count = 0
    for raw_fact in extracted.data.get('facts', []):
        statement = clean_text(raw_fact.get('statement', ''))
        title = clean_text(raw_fact.get('title', ''))[:220] or statement[:80] or 'Profile fact'
        fact_type = clean_text(raw_fact.get('fact_type', 'achievement')).lower().replace(' ', '_')
        if fact_type not in {choice[0] for choice in ProfileFact.FACT_CHOICES}:
            fact_type = 'achievement'
        key = (fact_type, title, statement)
        if not statement or key in existing_keys:
            continue
        evidence_quote = clean_text(raw_fact.get('evidence_quote', '')) or statement
        source_chunk = next((chunk for chunk in chunks if evidence_quote[:80].lower() in chunk.text.lower()), chunks[0] if chunks else None)
        ambiguous = bool(raw_fact.get('ambiguous')) and ambiguity_count < 5
        if ambiguous:
            ambiguity_count += 1
        fact = ProfileFact.objects.create(
            owner=document.owner,
            fact_type=fact_type,
            title=title,
            statement=statement,
            confidence=clean_text(raw_fact.get('confidence', 'medium'))[:24] or 'medium',
            source_document=document,
            source_chunk=source_chunk,
            metadata={
                'extractor': extracted.source,
                'onboarding_ambiguous': ambiguous,
                'ambiguity_reason': clean_text(raw_fact.get('ambiguity_reason', '')) if ambiguous else '',
            },
            lifecycle='proposed',
            evidence_quote=evidence_quote,
            strength='strong' if raw_fact.get('confidence') == 'high' and not ambiguous else 'working',
        )
        from .domain.embeddings import refresh_fact_embedding

        refresh_fact_embedding(fact)
        existing_keys.add(key)
        created += 1

    document.status = 'ready'
    document.status_message = f'Created {created} facts from {len(chunks)} chunks.'
    document.metadata = {
        **(document.metadata or {}),
        'extractor': extracted.source,
        'resume_analysis': {
            'overview': clean_text(extracted.data.get('overview', '')),
            'career_headline': clean_text(extracted.data.get('career_headline', '')),
            'likely_location': clean_text(extracted.data.get('likely_location', '')),
            'likely_industries': [clean_text(value) for value in extracted.data.get('likely_industries', []) if clean_text(value)][:12],
            'ambiguities': ambiguity_count,
        } if document.kind == 'resume' else {},
    }
    document.save(update_fields=['status', 'status_message', 'metadata', 'updated_at'])

    if document.kind == 'resume':
        ensure_canonical_resume(document)

    from .domain.profiles import candidate_profile, compute_profile_completeness

    if document.kind == 'resume':
        profile = candidate_profile(document.owner)
        profile = CandidateProfile.objects.select_for_update().get(pk=profile.pk)
        state = dict(profile.onboarding_state or {})
        state.pop('current_question', None)
        state['resume_document_id'] = document.id
        state['resume_analyzed_at'] = timezone.now().isoformat()
        profile.onboarding_state = state
        profile.save(update_fields=['onboarding_state', 'updated_at'])

    compute_profile_completeness(document.owner)
    from .domain.embeddings import refresh_profile_embedding

    refresh_profile_embedding(document.owner)

    return {'created_facts': created, 'created_chunks': len(chunks)}


def ensure_canonical_resume(document: ProfileDocument) -> Resume:
    existing = Resume.objects.filter(owner=document.owner, kind='canonical').order_by('-updated_at').first()
    title = 'Canonical Resume'
    content = document.raw_text.strip()
    if existing:
        existing.title = title
        existing.content_markdown = content
        existing.content_json = {'source_document_id': document.id}
        existing.save(update_fields=['title', 'content_markdown', 'content_json', 'updated_at'])
        return existing
    return Resume.objects.create(
        owner=document.owner,
        kind='canonical',
        title=title,
        content_markdown=content,
        content_json={'source_document_id': document.id},
    )


@transaction.atomic
def import_job_posting(owner, *, text: str, source_url: str = '', source=None, score: bool = True) -> JobPosting:
    extracted = extract_job(text, source_url=source_url)
    data = extracted.data
    content_hash = stable_hash(f'{owner.pk}:{source_url}:{text}')
    defaults = {
        'source': source,
        'title': clean_text(data.get('title'))[:240] or 'Imported Job',
        'company': clean_text(data.get('company'))[:220],
        'location': clean_text(data.get('location'))[:220],
        'remote_policy': normalize_remote_policy(data.get('remote_policy')),
        'seniority': clean_text(data.get('seniority'))[:120],
        'compensation': clean_text(data.get('compensation'))[:160],
        'description_text': text,
        'extracted_json': {**data, 'extractor': extracted.source},
        'source_url': source_url,
        'application_url': clean_text(data.get('application_url'))[:1000] or source_url,
    }
    job, _ = JobPosting.objects.update_or_create(
        owner=owner,
        content_hash=content_hash,
        defaults=defaults,
    )
    from .domain.sourcing import persist_job_structure

    job.canonical_url = source_url
    job.last_seen_at = timezone.now()
    job.freshness_status = 'fresh'
    job.save(update_fields=['canonical_url', 'last_seen_at', 'freshness_status', 'updated_at'])
    persist_job_structure(job)
    if score:
        from .domain.embeddings import refresh_job_embedding

        refresh_job_embedding(job)
        recompute_match(job)
    return job


def normalize_remote_policy(value: Any) -> str:
    lowered = clean_text(value).lower()
    if 'remote' in lowered:
        return 'remote'
    if 'hybrid' in lowered:
        return 'hybrid'
    if 'site' in lowered or 'office' in lowered:
        return 'onsite'
    return 'unknown'


def recompute_match(job: JobPosting) -> JobMatch:
    from .domain.matching import recompute_match as recompute_match_v2

    return recompute_match_v2(job)


def build_match_summary(score: int, covered: list[str], missing: list[str]) -> str:
    if score >= 80:
        prefix = 'Strong match'
    elif score >= 55:
        prefix = 'Possible match'
    else:
        prefix = 'Weak match'
    coverage = f'{len(covered)} covered skills'
    gap = f'{len(missing)} visible gaps' if missing else 'no obvious skill gaps'
    return f'{prefix}: {coverage}, {gap}.'


def create_tailored_resume(owner, *, job: JobPosting, canonical: Resume | None = None) -> Resume:
    if canonical is None:
        canonical = Resume.objects.filter(owner=owner, kind='canonical').order_by('-updated_at').first()
    if canonical is None:
        canonical = Resume.objects.create(
            owner=owner,
            kind='canonical',
            title='Canonical Resume',
            content_markdown='',
            content_json={},
        )
    facts = list(ProfileFact.objects.filter(owner=owner).order_by('-verified_by_user', 'fact_type', 'title')[:120])
    profile = getattr(owner, 'candidate_profile', None)
    result = tailor_resume(
        canonical_markdown=canonical.content_markdown,
        job_title=job.title,
        job_text=job.description_text,
        facts=[{
            'id': fact.id,
            'title': fact.title,
            'statement': fact.statement,
            'verified_by_user': fact.verified_by_user,
        } for fact in facts],
        candidate_name=owner.get_full_name() or owner.get_username().replace('.', ' ').replace('_', ' ').title(),
        candidate_headline=getattr(profile, 'headline', ''),
        candidate_location=getattr(profile, 'location', ''),
    )
    validation = {
        'generator': result.source,
        'summary_changes': result.data.get('summary_changes', []),
        'keyword_coverage': result.data.get('keyword_coverage', []),
        'unsupported_claims': result.data.get('unsupported_claims', []),
        'weak_claims': result.data.get('weak_claims', []),
        'evidence_links': result.data.get('evidence_links', []),
        'risk_notes': result.data.get('risk_notes', []),
    }
    return Resume.objects.create(
        owner=owner,
        kind='tailored',
        title=clean_text(result.data.get('title'))[:220] or f'{job.title} Tailored Resume',
        content_markdown=result.data.get('content_markdown', canonical.content_markdown),
        content_json={
            'target_job_id': job.id,
            'design': result.data.get('design') or {
                'template': 'modern', 'density': 'balanced', 'accent': '#177d69',
                'page_size': 'Letter', 'rationale': 'A clean, ATS-safe presentation selected for this application.',
            },
        },
        parent_resume=canonical,
        target_job=job,
        validation=validation,
    )


def generate_strategy(owner) -> dict[str, Any]:
    apps = Application.objects.filter(owner=owner)
    total = apps.count()
    by_status = dict(apps.values_list('status').annotate(count=Count('id')).values_list('status', 'count'))
    interviews = sum(by_status.get(status, 0) for status in ['recruiter_screen', 'technical_screen', 'onsite_final', 'offer'])
    applied = by_status.get('applied', 0) + interviews + by_status.get('rejected', 0)
    response_rate = round((interviews / applied) * 100, 1) if applied else 0
    followups_due = apps.filter(
        follow_up_at__lte=timezone.now(),
    ).exclude(status__in=['rejected', 'archived', 'offer']).count()
    top_matches = JobMatch.objects.filter(owner=owner).select_related('job').order_by('-score')[:5]

    recommendations = []
    if top_matches:
        best = top_matches[0]
        recommendations.append({
            'title': f'Prioritize {best.job.title}',
            'detail': f'This is currently the highest-scoring match at {best.score}.',
        })
    if followups_due:
        recommendations.append({
            'title': f'Follow up on {followups_due} application{"s" if followups_due != 1 else ""}',
            'detail': 'These applications have follow-up dates due or overdue.',
        })
    weak_facts = ProfileFact.objects.filter(owner=owner, verified_by_user=False).count()
    if weak_facts:
        recommendations.append({
            'title': 'Verify profile facts',
            'detail': f'{weak_facts} extracted facts are unverified and may limit resume confidence.',
        })
    if not recommendations:
        recommendations.append({
            'title': 'Import more jobs',
            'detail': 'Add job descriptions to improve matching and strategy feedback.',
        })

    return {
        'totals': {
            'applications': total,
            'applied': applied,
            'interviews': interviews,
            'response_rate': response_rate,
            'followups_due': followups_due,
        },
        'by_status': by_status,
        'top_matches': [
            {'job_id': match.job_id, 'title': match.job.title, 'company': match.job.company, 'score': match.score}
            for match in top_matches
        ],
        'recommendations': recommendations,
    }


def dashboard(owner) -> dict[str, Any]:
    strategy = generate_strategy(owner)
    return {
        'profile_documents': ProfileDocument.objects.filter(owner=owner).count(),
        'profile_facts': ProfileFact.objects.filter(owner=owner).count(),
        'jobs': JobPosting.objects.filter(owner=owner).count(),
        'matches': JobMatch.objects.filter(owner=owner).count(),
        'applications': Application.objects.filter(owner=owner).count(),
        'resumes': Resume.objects.filter(owner=owner).count(),
        'strategy': strategy,
    }
