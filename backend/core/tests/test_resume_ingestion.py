from pathlib import Path
from tempfile import TemporaryDirectory
from unittest.mock import patch

from django.contrib.auth import get_user_model
from django.test import TestCase, override_settings
from docx import Document
from reportlab.pdfgen import canvas

from core.ai import AIResult
from core.domain.profiles import answer_onboarding, onboarding_snapshot
from core.models import CandidateProfile, OnboardingResponse, ProfileDocument, ProfileFact
from core.services import extract_text_from_upload, ingest_profile_document


class ResumeFormatTests(TestCase):
    def test_extracts_pdf_docx_html_rtf_and_plain_text(self):
        with TemporaryDirectory() as directory:
            root = Path(directory)
            pdf = root / 'resume.pdf'
            pdf_canvas = canvas.Canvas(str(pdf))
            pdf_canvas.drawString(72, 720, 'PDF resume: led a platform reliability program')
            pdf_canvas.save()

            docx = root / 'resume.docx'
            word = Document()
            word.add_heading('Word Resume', 0)
            word.add_paragraph('Built dependable Django services.')
            table = word.add_table(rows=1, cols=1)
            table.cell(0, 0).text = 'Python and PostgreSQL'
            word.save(docx)

            html = root / 'resume.html'
            html.write_text('<h1>HTML Resume</h1><p>Shipped an AI platform.</p><script>ignore me</script>', encoding='utf-8')
            rtf = root / 'resume.rtf'
            rtf.write_text(r'{\rtf1\ansi RTF Resume\par Built APIs with Redis.}', encoding='utf-8')
            plain = root / 'resume.md'
            plain.write_text('# Markdown Resume\n\nLed engineering teams.', encoding='utf-8')

            self.assertIn('platform reliability', extract_text_from_upload(str(pdf)))
            self.assertIn('Python and PostgreSQL', extract_text_from_upload(str(docx)))
            self.assertIn('HTML Resume', extract_text_from_upload(str(html)))
            self.assertNotIn('ignore me', extract_text_from_upload(str(html)))
            self.assertIn('Built APIs with Redis', extract_text_from_upload(str(rtf)))
            self.assertIn('Led engineering teams', extract_text_from_upload(str(plain)))


class AdaptiveResumeInterviewTests(TestCase):
    @override_settings(OPENAI_API_KEY='')
    def test_legacy_saved_answers_are_replayed_once_instead_of_asked_again(self):
        user = get_user_model().objects.create_user(username='resume-existing-answers', password='password')
        ProfileDocument.objects.create(
            owner=user, kind='resume', title='resume.txt', raw_text='Existing resume', status='ready',
        )
        profile = CandidateProfile.objects.create(owner=user, onboarding_state={
            'started': True,
            'answered_targets': ['employment_types', 'professional_summary'],
            'interview_history': [
                {'question_id': 'q1-employment_types-0', 'target': 'employment_types', 'question': 'Employment', 'answer': 'Full-time'},
                {'question_id': 'q2-professional_summary-0', 'target': 'professional_summary', 'question': 'Summary', 'answer': 'Platform leader who builds dependable systems and helps teams deliver measurable outcomes.'},
            ],
        })

        onboarding_snapshot(user)

        profile.refresh_from_db()
        self.assertEqual(profile.employment_types, ['full-time'])
        self.assertIn('dependable systems', profile.professional_summary)
        self.assertIn('legacy_answers_reconciled_at', profile.onboarding_state)
        self.assertEqual(OnboardingResponse.objects.filter(owner=user).count(), 2)

    @override_settings(OPENAI_API_KEY='')
    @patch('core.services.analyze_candidate_resume')
    def test_ambiguous_resume_claim_is_confirmed_before_other_questions(self, analyze):
        user = get_user_model().objects.create_user(username='dynamic', password='password')
        analyze.return_value = AIResult(data={
            'overview': 'A platform engineering resume with one unclear ownership claim.',
            'career_headline': 'Platform Engineer',
            'likely_location': 'Toronto, Canada',
            'likely_industries': ['Developer tools'],
            'facts': [{
                'fact_type': 'achievement',
                'title': 'Revenue impact',
                'statement': 'Increased revenue by 30 percent.',
                'evidence_quote': 'Contributed to initiatives that increased revenue 30%.',
                'confidence': 'low',
                'ambiguous': True,
                'ambiguity_reason': 'The resume does not make the candidate’s personal contribution clear.',
            }],
        }, source='openai')
        document = ProfileDocument.objects.create(
            owner=user, kind='resume', title='current.html',
            raw_text='Contributed to initiatives that increased revenue 30%.',
        )
        ingest_profile_document(document)
        profile = user.candidate_profile
        profile.onboarding_state = {**profile.onboarding_state, 'started': True}
        profile.save(update_fields=['onboarding_state', 'updated_at'])

        snapshot = onboarding_snapshot(user)
        question = snapshot['step']['question']
        self.assertEqual(question['target'], 'fact_confirmation')
        self.assertIn('Contributed to initiatives', question['evidence'])

        corrected = 'Contributed backend delivery to a broader initiative that increased company revenue by 30 percent.'
        answer_onboarding(user, step='interview', answers={
            'question_id': question['id'],
            'value': corrected,
        })
        fact = ProfileFact.objects.get(owner=user, title='Revenue impact')
        self.assertTrue(fact.verified_by_user)
        self.assertEqual(fact.lifecycle, 'verified')
        self.assertEqual(fact.statement, corrected)
        self.assertFalse(fact.metadata['onboarding_ambiguous'])
